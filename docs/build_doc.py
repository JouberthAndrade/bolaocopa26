# -*- coding: utf-8 -*-
"""Gera o Projeto Integrador (BolaoCopa2026) em .docx com formatacao ABNT."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
IMG_CLASSES = os.path.join(BASE, "diagramas", "diagrama_classes.png")
IMG_DER = os.path.join(BASE, "diagramas", "diagrama_der.png")
OUT = os.path.join(BASE, "Projeto_Integrador_BolaoCopa2026.docx")

FONT = "Times New Roman"
doc = Document()

# ---------- margens ABNT (esq/sup 3cm, dir/inf 2cm) ----------
sec = doc.sections[0]
sec.top_margin = Cm(3)
sec.left_margin = Cm(3)
sec.right_margin = Cm(2)
sec.bottom_margin = Cm(2)

# ---------- estilo base ----------
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def _set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)


def body(text, justify=True, indent=True, spacing=1.5, size=12, space_after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE if spacing == 1.5 else WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if indent:
        pf.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    _set_font(r, size=size)
    return p


def heading(num, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(18 if level == 1 else 12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.keep_with_next = True
    label = f"{num} {text}" if num else text
    r = p.add_run(label.upper() if level == 1 else label)
    _set_font(r, size=12, bold=True)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(3)
    pf.left_indent = Cm(1.25)
    r = p.add_run(text)
    _set_font(r, size=12)
    return p


def centered(text, size=12, bold=False, space_after=0, caps=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text.upper() if caps else text)
    _set_font(r, size=size, bold=bold, italic=italic)
    return p


def blank(n=1):
    for _ in range(n):
        centered("")


# ============================================================
# CAPA
# ============================================================
centered("UNIVERSIDADE SANTO AMARO", size=12, bold=True, space_after=0)
centered("CURSO DE CIÊNCIA DA COMPUTAÇÃO", size=12, bold=True)
blank(6)
centered("JOUBERTH ANDRADE", size=12, bold=True)
blank(6)
centered(
    "BOLÃO COPA DO MUNDO FIFA 2026: DESENVOLVIMENTO DE UMA PLATAFORMA WEB "
    "DE PALPITES ESPORTIVOS UTILIZANDO NEXT.JS, TYPESCRIPT E ARQUITETURA EM CAMADAS",
    size=12, bold=True,
)
blank(12)
centered("Belo Horizonte", size=12)
centered("2026", size=12)
doc.add_page_break()

# ============================================================
# FOLHA DE ROSTO
# ============================================================
centered("JOUBERTH ANDRADE", size=12, bold=True)
blank(8)
centered(
    "BOLÃO COPA DO MUNDO FIFA 2026: DESENVOLVIMENTO DE UMA PLATAFORMA WEB "
    "DE PALPITES ESPORTIVOS UTILIZANDO NEXT.JS, TYPESCRIPT E ARQUITETURA EM CAMADAS",
    size=12, bold=True,
)
blank(4)
# bloco natureza do trabalho (recuado a partir do meio da pagina)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(8)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
r = p.add_run(
    "Projeto Integrador apresentado ao Curso de Ciência da Computação da "
    "Universidade Santo Amaro como requisito parcial para aprovação na "
    "disciplina Projeto Integrador: Desenvolvimento de Sistemas."
)
_set_font(r, size=10)
blank(1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(8)
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
r = p.add_run("Orientador: Prof. Angelo Luiz da Cruz Oliveira")
_set_font(r, size=10)
blank(10)
centered("Belo Horizonte", size=12)
centered("2026", size=12)
doc.add_page_break()

# ============================================================
# SUMARIO
# ============================================================
centered("SUMÁRIO", size=12, bold=True, space_after=12)
sumario = [
    ("1 INTRODUÇÃO", "4"),
    ("2 PROBLEMA", "5"),
    ("3 HIPÓTESE", "5"),
    ("4 OBJETIVOS", "6"),
    ("4.1 Objetivo geral", "6"),
    ("4.2 Objetivos específicos", "6"),
    ("5 JUSTIFICATIVA", "7"),
    ("6 REFERENCIAL TEÓRICO", "8"),
    ("6.1 Bolões esportivos e gamificação", "8"),
    ("6.2 Arquitetura web moderna: Next.js e React Server Components", "9"),
    ("6.3 Persistência de dados, ORM e o modelo relacional", "9"),
    ("6.4 Princípios SOLID e segurança de aplicações web", "10"),
    ("7 MODELAGEM DO SISTEMA", "11"),
    ("7.1 Diagrama de classes (UML)", "11"),
    ("7.2 Diagrama entidade-relacionamento (DER)", "12"),
    ("8 METODOLOGIA", "13"),
    ("9 CRONOGRAMA DE ATIVIDADES", "14"),
    ("REFERÊNCIAS", "15"),
]
for label, page in sumario:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    # tab com leader de pontos
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT)
    pPr = p._p.get_or_add_pPr()
    tabsxml = pPr.find(qn("w:tabs"))
    if tabsxml is not None:
        tabsxml[-1].set(qn("w:leader"), "dot")
    bold = "." not in label.split(" ")[0]
    r = p.add_run(label + "\t" + page)
    _set_font(r, size=12, bold=bold)
doc.add_page_break()

# ============================================================
# 1 INTRODUCAO
# ============================================================
heading("1", "Introdução")
body("As competições esportivas de grande alcance, como a Copa do Mundo FIFA, "
     "mobilizam bilhões de pessoas e estimulam práticas sociais de confraternização, "
     "entre as quais se destacam os chamados bolões — apostas informais nas quais "
     "grupos de amigos, familiares ou colegas de trabalho registram seus palpites "
     "sobre os resultados das partidas e disputam uma classificação ao longo do "
     "torneio. Tradicionalmente conduzidos em papel ou em planilhas, esses bolões "
     "apresentam limitações relacionadas à apuração manual de pontos, à "
     "transparência das regras e ao acompanhamento da classificação em tempo real.")
body("No contexto acadêmico, a construção de sistemas que digitalizam práticas "
     "cotidianas constitui uma estratégia eficaz para a consolidação de "
     "conhecimentos em programação e engenharia de software. Conforme destacam "
     "Pressman e Maxim (2021), o desenvolvimento de software por meio de projetos "
     "aplicados permite ao estudante exercitar não apenas habilidades técnicas de "
     "codificação, mas também competências de projeto de arquitetura, organização "
     "de código e aplicação de padrões de desenvolvimento.")
body("Diante desse cenário, o presente trabalho propõe o desenvolvimento de uma "
     "plataforma web de bolões para a Copa do Mundo FIFA 2026, na qual os usuários "
     "podem criar bolões privados, convidar participantes, registrar palpites para "
     "as partidas e acompanhar um ranking atualizado de forma automática. O sistema "
     "é desenvolvido com o framework Next.js (React) e a linguagem TypeScript, "
     "utilizando o ORM Prisma sobre um banco de dados relacional PostgreSQL, "
     "organizado segundo uma arquitetura em camadas que separa apresentação, "
     "regras de negócio e acesso a dados.")
body("O tema permite integrar diversos conceitos estudados ao longo da disciplina, "
     "tais como: modelagem de dados, orientação a objetos, validação de entradas, "
     "comunicação cliente-servidor, autenticação de usuários e organização modular "
     "do código-fonte. Como contribuição adicional, este projeto apresenta a "
     "modelagem formal do sistema por meio de um diagrama de classes em UML "
     "(Unified Modeling Language) e de um diagrama entidade-relacionamento (DER) "
     "do banco de dados, recursos que documentam a estrutura da solução e "
     "facilitam sua compreensão e manutenção.")

# ============================================================
# 2 PROBLEMA
# ============================================================
heading("2", "Problema")
body("Como desenvolver uma plataforma web funcional para a gestão de bolões "
     "esportivos, capaz de registrar palpites, apurar pontuações de forma "
     "automática e apresentar a classificação dos participantes, empregando boas "
     "práticas de engenharia de software e de modelagem de dados?")
body("A formulação deste problema parte da observação de que muitos projetos "
     "acadêmicos de programação limitam-se a exercícios desconectados de aplicações "
     "reais, o que dificulta a assimilação de conceitos como separação de "
     "responsabilidades, modelagem de dados e comunicação entre camadas de "
     "software. A proposta de digitalizar uma prática amplamente conhecida — o "
     "bolão esportivo — oferece um contexto rico para a aplicação integrada desses "
     "conhecimentos.")

# ============================================================
# 3 HIPOTESE
# ============================================================
heading("3", "Hipótese")
body("A utilização de uma arquitetura web em camadas, combinada com um banco de "
     "dados relacional modelado a partir de um diagrama entidade-relacionamento e "
     "com a aplicação dos princípios SOLID, permite a construção de uma plataforma "
     "de bolões modular, segura e de fácil manutenção, capaz de automatizar a "
     "apuração de pontos e o cálculo da classificação dos participantes.")

# ============================================================
# 4 OBJETIVOS
# ============================================================
heading("4", "Objetivos")
heading("4.1", "Objetivo geral", level=2)
body("Desenvolver uma plataforma web para a criação e a gestão de bolões da Copa "
     "do Mundo FIFA 2026, contemplando o registro de palpites, a apuração "
     "automática de pontos e a exibição da classificação dos participantes, "
     "utilizando Next.js, TypeScript e arquitetura em camadas.")
heading("4.2", "Objetivos específicos", level=2)
for it in [
    "Modelar o domínio do sistema por meio de um diagrama de classes em UML;",
    "Projetar o banco de dados relacional por meio de um diagrama "
    "entidade-relacionamento (DER);",
    "Implementar a autenticação de usuários e o controle de acesso aos bolões;",
    "Desenvolver o registro de palpites com validação de prazo e de integridade "
    "dos dados;",
    "Implementar a apuração automática de pontos a partir dos resultados das "
    "partidas e das regras de pontuação configuráveis;",
    "Disponibilizar a classificação (ranking) dos participantes de cada bolão;",
    "Aplicar os princípios SOLID e padrões de projeto na estruturação do "
    "código-fonte.",
]:
    bullet(it)

# ============================================================
# 5 JUSTIFICATIVA
# ============================================================
heading("5", "Justificativa")
body("A justificativa para o desenvolvimento deste projeto reside na importância "
     "de integrar os conhecimentos adquiridos ao longo da disciplina, aplicando-os "
     "em um contexto prático e significativo. Segundo Sommerville (2019), o "
     "desenvolvimento de projetos de software aplicados contribui para a formação "
     "de profissionais capazes de lidar com os desafios reais da indústria de "
     "tecnologia.")
body("A temática dos bolões esportivos possibilita a exploração de conceitos "
     "fundamentais da computação, tais como: modelagem de dados, validação de "
     "entradas, persistência em banco relacional, separação de responsabilidades e "
     "comunicação cliente-servidor. De acordo com Kahlmeyer-Mertens et al. (2007, "
     "p. 50), “como o nome já diz, trata-se da ação de justificar ou ainda dar "
     "justiça à escolha do tema e do trabalho”.")
body("Além disso, a escolha de uma arquitetura web em camadas permite que o "
     "sistema seja facilmente estendido no futuro com novos recursos, demonstrando "
     "a aplicabilidade dos conceitos de desacoplamento entre as partes da "
     "aplicação. A utilização dos princípios SOLID, conforme propostos por Martin "
     "(2017), garante que o código produzido seja de alta qualidade, facilitando a "
     "manutenção e a evolução do sistema.")
body("O projeto também se justifica pela relevância social e cultural do tema, "
     "uma vez que os bolões da Copa do Mundo são uma prática de grande alcance no "
     "Brasil. Digitalizar essa prática, tornando-a mais transparente e acessível, "
     "contribui para a aproximação entre tecnologia e cotidiano dos usuários.")

# ============================================================
# 6 REFERENCIAL TEORICO
# ============================================================
heading("6", "Referencial Teórico")
heading("6.1", "Bolões esportivos e gamificação", level=2)
body("Os bolões esportivos são competições informais nas quais os participantes "
     "registram previsões sobre os resultados de eventos esportivos e acumulam "
     "pontos conforme a precisão de seus palpites. Sob a ótica da engenharia de "
     "software, transpor essa prática para um sistema digital exige a definição "
     "clara de regras de pontuação, de prazos para o registro de palpites e de "
     "critérios de classificação. Segundo Deterding et al. (2011), a gamificação — "
     "o uso de elementos de jogos em contextos não lúdicos, como rankings, pontos "
     "e recompensas — aumenta o engajamento dos usuários, o que justifica a "
     "presença de mecanismos de classificação e de feedback contínuo na plataforma.")
heading("6.2", "Arquitetura web moderna: Next.js e React Server Components", level=2)
body("O React é uma biblioteca para a construção de interfaces de usuário baseada "
     "em componentes reutilizáveis. O Next.js, por sua vez, é um framework "
     "construído sobre o React que adiciona renderização no servidor, roteamento "
     "baseado em arquivos e a execução de lógica de servidor por meio de Server "
     "Actions e Route Handlers (VERCEL, 2024). Essa abordagem favorece a separação "
     "entre a camada de apresentação, executada no navegador, e a camada de "
     "negócio, executada no servidor, contribuindo para a segurança e o desempenho "
     "da aplicação. No presente projeto, regras sensíveis — como o travamento de "
     "palpites antes do início das partidas — são validadas no servidor, evitando "
     "que sejam contornadas pelo cliente.")
heading("6.3", "Persistência de dados, ORM e o modelo relacional", level=2)
body("O modelo relacional, proposto por Codd (1970) e consolidado por Date (2004), "
     "organiza os dados em tabelas (relações) compostas por linhas e colunas, "
     "relacionadas por meio de chaves primárias e estrangeiras. O diagrama "
     "entidade-relacionamento (DER), descrito por Heuser (2009), é uma técnica de "
     "modelagem que representa graficamente entidades, atributos e os "
     "relacionamentos entre elas, com suas respectivas cardinalidades. Para mediar "
     "a comunicação entre o código orientado a objetos e o banco de dados "
     "relacional, utiliza-se um ORM (Object-Relational Mapping). O Prisma é um ORM "
     "moderno para o ecossistema JavaScript/TypeScript que define o esquema do "
     "banco de forma declarativa e gera um cliente fortemente tipado, reduzindo "
     "erros e aumentando a produtividade (PRISMA, 2024). O banco de dados adotado é "
     "o PostgreSQL, um sistema gerenciador de banco de dados relacional de código "
     "aberto reconhecido por sua robustez e conformidade com o padrão SQL "
     "(POSTGRESQL, 2024).")
heading("6.4", "Princípios SOLID e segurança de aplicações web", level=2)
body("Os princípios SOLID, propostos por Robert C. Martin (2017), representam um "
     "conjunto de cinco diretrizes para o projeto de software orientado a objetos: "
     "Responsabilidade Única (SRP), Aberto-Fechado (OCP), Substituição de Liskov "
     "(LSP), Segregação de Interfaces (ISP) e Inversão de Dependência (DIP). Esses "
     "princípios visam à produção de código limpo, desacoplado e de fácil "
     "manutenção. Segundo Fowler (2002), a separação em camadas e a abstração do "
     "acesso a dados permitem que mudanças na infraestrutura de persistência não "
     "afetem as regras de negócio. No que se refere à segurança, adotam-se "
     "práticas como o armazenamento de senhas por meio de funções de hash, a "
     "validação de todas as entradas do usuário e o controle de acesso baseado na "
     "associação do usuário a cada bolão.")

# ============================================================
# 7 MODELAGEM DO SISTEMA
# ============================================================
heading("7", "Modelagem do Sistema")
body("Esta seção apresenta a modelagem da plataforma de bolões sob duas "
     "perspectivas complementares: a estrutura orientada a objetos do domínio, "
     "representada por um diagrama de classes em UML, e a estrutura de "
     "persistência, representada por um diagrama entidade-relacionamento (DER) do "
     "banco de dados relacional. Ambos os diagramas foram elaborados a partir do "
     "código-fonte e do esquema de dados efetivamente implementados no sistema.")

heading("7.1", "Diagrama de classes (UML)", level=2)
body("O diagrama de classes descreve as principais entidades do domínio, seus "
     "atributos, operações e relacionamentos. Destacam-se a entidade Usuário "
     "(User), que representa tanto participantes humanos quanto palpiteiros "
     "automatizados de inteligência artificial; o Bolão (Pool), que agrega os "
     "participantes por meio de associações (Membership); o Palpite (Bet), que "
     "registra a previsão de um usuário para uma partida; e a Partida (Match), "
     "vinculada a duas seleções (Team) e a um Torneio (Tournament). As classes de "
     "serviço — ScoringService, RankingService e SyncService — concentram as regras "
     "de negócio de pontuação, classificação e sincronização de resultados, "
     "evidenciando a separação entre o modelo de domínio e a lógica de aplicação. "
     "A Figura 1 apresenta o diagrama de classes do sistema.")
doc.add_picture(IMG_CLASSES, width=Cm(16))
fig = doc.paragraphs[-1]
fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
centered("Figura 1 – Diagrama de classes (UML) da plataforma de bolões",
         size=10, space_after=0)
centered("Fonte: o autor (2026).", size=10, space_after=12)

heading("7.2", "Diagrama entidade-relacionamento (DER)", level=2)
body("O diagrama entidade-relacionamento representa as tabelas do banco de dados "
     "PostgreSQL, seus atributos, chaves primárias (PK), chaves estrangeiras (FK) "
     "e restrições de unicidade (UQ), bem como as cardinalidades dos "
     "relacionamentos na notação pé-de-galinha (crow's foot). O modelo evidencia, "
     "por exemplo, que um usuário pode participar de vários bolões e que cada "
     "bolão possui exatamente uma regra de pontuação (ScoringRule). As entidades "
     "de autenticação (Account e Session) decorrem da integração com a biblioteca "
     "de autenticação adotada. A Figura 2 apresenta o DER do sistema.")
doc.add_picture(IMG_DER, width=Cm(16))
fig = doc.paragraphs[-1]
fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
centered("Figura 2 – Diagrama entidade-relacionamento (DER) do banco de dados",
         size=10, space_after=0)
centered("Fonte: o autor (2026).", size=10, space_after=12)

# ============================================================
# 8 METODOLOGIA
# ============================================================
heading("8", "Metodologia")
body("O presente trabalho caracteriza-se como uma pesquisa aplicada de natureza "
     "qualitativa, com abordagem exploratória, cujo delineamento se configura como "
     "um estudo de caso de desenvolvimento de software. Segundo Gil (2010), a "
     "pesquisa aplicada tem como finalidade gerar conhecimentos para aplicação "
     "prática, dirigidos à solução de problemas específicos.")
body("O percurso metodológico foi organizado nas seguintes etapas:")
for it in [
    "Levantamento dos requisitos do sistema e definição das regras de negócio do "
    "bolão (registro de palpites, prazos, pontuação e classificação);",
    "Modelagem do sistema por meio do diagrama de classes (UML) e do diagrama "
    "entidade-relacionamento (DER);",
    "Definição da arquitetura em camadas — apresentação (componentes e páginas), "
    "negócio (services) e acesso a dados (ORM/repositórios);",
    "Implementação do backend com Next.js, TypeScript e Prisma, incluindo "
    "autenticação, Server Actions e serviços de pontuação e ranking;",
    "Implementação do frontend com componentes React e estilização com Tailwind "
    "CSS;",
    "Integração com fonte externa de resultados das partidas e automação da "
    "sincronização e da apuração por tarefas agendadas;",
    "Testes das regras de negócio (pontuação, travamento de palpites e ranking) "
    "e documentação do projeto.",
]:
    bullet(it)
body("O ambiente de desenvolvimento foi composto por: Node.js 24; framework "
     "Next.js 15 (App Router); linguagem TypeScript; ORM Prisma sobre banco de "
     "dados PostgreSQL; biblioteca de autenticação Auth.js v5; framework de "
     "estilização Tailwind CSS; e o editor Visual Studio Code. A hospedagem é "
     "realizada na plataforma Vercel, com tarefas agendadas (cron) responsáveis "
     "pela sincronização periódica dos jogos e pela apuração dos pontos. Os "
     "diagramas de classes e entidade-relacionamento foram gerados com a "
     "ferramenta PlantUML a partir do modelo implementado.")

# ============================================================
# 9 CRONOGRAMA
# ============================================================
heading("9", "Cronograma de Atividades")
body("De acordo com Alves (2007), o cronograma é a organização das fases e etapas "
     "do projeto. O Quadro 1 apresenta o cronograma de atividades do presente "
     "trabalho.")
centered("Quadro 1 – Cronograma de atividades", size=10, space_after=4)

atividades = [
    ("Atividades", "Fev", "Mar", "Abr", "Mai", "Jun"),
    ("Levantamento de requisitos", "X", "", "", "", ""),
    ("Modelagem (UML e DER)", "X", "X", "", "", ""),
    ("Definição da arquitetura", "", "X", "", "", ""),
    ("Implementação do backend", "", "X", "X", "", ""),
    ("Implementação do frontend", "", "", "X", "X", ""),
    ("Testes e validação", "", "", "", "X", ""),
    ("Documentação do projeto", "", "", "", "X", "X"),
    ("Revisão final e apresentação", "", "", "", "", "X"),
]
tbl = doc.add_table(rows=len(atividades), cols=6)
tbl.style = "Table Grid"
tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, row in enumerate(atividades):
    for j, val in enumerate(row):
        cell = tbl.cell(i, j)
        cell.text = ""
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(0)
        rr = cp.add_run(val)
        _set_font(rr, size=10, bold=(i == 0))
centered("Fonte: o autor (2026).", size=10, space_after=12)

# ============================================================
# REFERENCIAS
# ============================================================
doc.add_page_break()
heading("", "Referências")

refs = [
    "ALVES, Magda. Como escrever teses e monografias: um roteiro passo a passo. "
    "2. ed. rev. atual. Rio de Janeiro: Campus, 2007.",
    "CODD, Edgar F. A relational model of data for large shared data banks. "
    "Communications of the ACM, v. 13, n. 6, p. 377-387, 1970.",
    "DATE, C. J. Introdução a sistemas de bancos de dados. 8. ed. Rio de Janeiro: "
    "Elsevier, 2004.",
    "DETERDING, Sebastian et al. From game design elements to gamefulness: "
    "defining gamification. In: PROCEEDINGS OF THE 15TH INTERNATIONAL ACADEMIC "
    "MINDTREK CONFERENCE. New York: ACM, 2011. p. 9-15.",
    "FOWLER, Martin. Patterns of enterprise application architecture. Boston: "
    "Addison-Wesley, 2002.",
    "GIL, Antonio Carlos. Como elaborar projetos de pesquisa. 5. ed. São Paulo: "
    "Atlas, 2010.",
    "HEUSER, Carlos Alberto. Projeto de banco de dados. 6. ed. Porto Alegre: "
    "Bookman, 2009.",
    "KAHLMEYER-MERTENS, Roberto S. et al. Como elaborar projetos de pesquisa: "
    "linguagem e método. Rio de Janeiro: FGV, 2007.",
    "MARTIN, Robert C. Arquitetura limpa: o guia do artesão para estrutura e "
    "design de software. Rio de Janeiro: Alta Books, 2017.",
    "POSTGRESQL GLOBAL DEVELOPMENT GROUP. PostgreSQL documentation. 2024. "
    "Disponível em: https://www.postgresql.org/docs/. Acesso em: 10 jun. 2026.",
    "PRESSMAN, Roger S.; MAXIM, Bruce R. Engenharia de software: uma abordagem "
    "profissional. 9. ed. Porto Alegre: AMGH, 2021.",
    "PRISMA. Prisma ORM documentation. 2024. Disponível em: "
    "https://www.prisma.io/docs. Acesso em: 10 jun. 2026.",
    "SOMMERVILLE, Ian. Engenharia de software. 10. ed. São Paulo: Pearson, 2019.",
    "VERCEL. Next.js documentation. 2024. Disponível em: https://nextjs.org/docs. "
    "Acesso em: 10 jun. 2026.",
]
for r in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(r)
    _set_font(run, size=12)

doc.save(OUT)
print("OK ->", OUT)
